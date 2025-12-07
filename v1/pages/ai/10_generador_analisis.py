"""
Generador de Análisis Inteligente - IA Agente
Crea scripts de análisis bajo demanda basado en peticiones del usuario
"""

import streamlit as st
import os
import sys
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv

# Cargar variables de entorno
load_dotenv('/home/cpinilla/projects/latam-3c/.env')

# Agregar path para imports
sys.path.insert(0, '/home/cpinilla/projects/latam-3c/v1')

from ai_modules.rag.rag_chain import RAGChain
from ai_modules.rag.sql_tool import SQLTool

# ============ CONFIGURACIÓN DE PÁGINA ============
st.set_page_config(
    page_title="🤖 Generador de Análisis",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============ ESTILOS ============
st.markdown("""
<style>
    .main-header {
        color: #1f77b4;
        font-size: 2.5rem;
        font-weight: bold;
        margin-bottom: 10px;
    }
    .section-header {
        color: #1f77b4;
        font-size: 1.3rem;
        font-weight: bold;
        margin-top: 20px;
        margin-bottom: 10px;
    }
    .status-box {
        padding: 15px;
        border-radius: 8px;
        margin: 10px 0;
        font-weight: 500;
    }
    .status-generating {
        background-color: #fff3cd;
        border-left: 4px solid #ffc107;
        color: #856404;
    }
    .status-success {
        background-color: #d4edda;
        border-left: 4px solid #28a745;
        color: #155724;
    }
    .status-error {
        background-color: #f8d7da;
        border-left: 4px solid #dc3545;
        color: #721c24;
    }
</style>
""", unsafe_allow_html=True)

# ============ FUNCIONES PRINCIPALES ============

def generate_analysis_script(peticion: str, rag_chain=None) -> str:
    """
    Genera un script Python de análisis.
    Intenta primero con RAGChain, fallback a Claude API directo.
    """
    import anthropic

    prompt = f"""Eres un especialista en análisis de datos de huella de carbono en la industria del cemento.

Genera un SCRIPT PYTHON que realice el análisis solicitado.

SINTAXIS OBLIGATORIA:
1. USA LISTA para construir el reporte: lineas = [] y lineas.append('texto')
2. Al final: reporte = '\\n'.join(lineas)
3. NUNCA uses += para strings largos
4. SQL en comillas simples: 'SELECT ...'
5. Máximo 200 líneas de código Python
6. Reporte final: 400-600 líneas máximo
7. Solo ASCII, sin emojis (usa -> en vez de →)

TÉCNICAS:
8. Conecta a PostgreSQL: psycopg2.connect(dbname='latam4c_db')
9. Guarda en /tmp/ con timestamp
10. try/except + conn.rollback() para queries
11. Columnas: "nombre_columna"

FORMATO REPORTE:
- CONCISO: Datos agregados (promedios, totales), NO listas completas
- ANALÍTICO: Interpretaciones + comparaciones con GCCA
- BENCHMARKING GCCA: Huella CO2 mundial 294 kg/m³, Europa 262 kg/m³
- ESTRUCTURA: Resumen + 3-5 secciones + conclusiones
- MÁXIMO: 400-600 líneas de reporte

TABLAS:
- huella_concretos: "año", "REST" (resistencia MPa), "volumen", "huella_co2", "origen", A1-A4_Total
- cementos: "origen", "planta", "cemento", "año", "huella_co2_bruta", "factor_clinker"
- plantas_latam: "planta", "ciudad", "pais", "compania", "grupo", lat, lon
- v_benchmarks_por_region: categoria, kpi_nombre, region, valor_promedio, valor_minimo, valor_maximo

JOINS:
huella_concretos."origen" = plantas_latam."planta"
⚠️ REST es resistencia MPa, NO es ID!

BENCHMARKING:
```sql
SELECT AVG(hc."huella_co2"), vbr.valor_promedio, vbr.region
FROM huella_concretos hc
CROSS JOIN v_benchmarks_por_region vbr
WHERE vbr.kpi_nombre = 'Energia Total de Produccion';
```

PETICIÓN DEL USUARIO:
{peticion}

Genera SOLO el código Python, sin explicaciones adicionales. El script debe ser funcional y ejecutable.

IMPORTANTE: USA LISTA para construir reporte (lineas.append() + '\\n'.join(lineas))

Ejemplo OBLIGATORIO:

```python
#!/usr/bin/env python3
import psycopg2
from datetime import datetime

# Conectar usando UNIX socket (sin password)
conn = psycopg2.connect(dbname='latam4c_db')
cursor = conn.cursor()

# ANÁLISIS CON DATOS AGREGADOS
try:
    cursor.execute('SELECT AVG("huella_co2"), MIN("huella_co2"), MAX("huella_co2"), COUNT(*) FROM huella_concretos')
    avg_co2, min_co2, max_co2, total = cursor.fetchone()
except Exception as e:
    conn.rollback()
    print("Error: " + str(e))

# Obtener referencias GCCA para comparación
try:
    cursor.execute('SELECT region, valor_promedio, valor_minimo, valor_maximo FROM v_benchmarks_por_region WHERE kpi_nombre = \\'Energía Total de Producción\\' ORDER BY valor_promedio')
    refs_gcca = cursor.fetchall()
except Exception as e:
    conn.rollback()
    refs_gcca = []

# Generar reporte CONCISO con benchmarking - USA LISTA
lineas = []
lineas.append("# Análisis de Huella de Carbono")
lineas.append("")
lineas.append("**Fecha:** " + datetime.now().strftime('%Y-%m-%d %H:%M'))
lineas.append("")
lineas.append("## Resumen Ejecutivo")
lineas.append("")
lineas.append("- Total registros: " + str(total))
lineas.append("- Huella CO2 promedio LATAM: " + str(round(avg_co2, 2)) + " kg CO2/m3")
lineas.append("- Rango: " + str(round(min_co2, 2)) + " - " + str(round(max_co2, 2)) + " kg CO2/m3")
lineas.append("")
lineas.append("## Benchmarking GCCA")
lineas.append("")
lineas.append("| Region | Promedio | Min | Max | vs LATAM |")
lineas.append("|--------|----------|-----|-----|----------|")
for region, prom, minval, maxval in refs_gcca:
    diff = round(avg_co2 - prom, 2)
    signo = "+" if diff > 0 else ""
    lineas.append("| " + region + " | " + str(prom) + " | " + str(minval) + " | " + str(maxval) + " | " + signo + str(diff) + " |")
lineas.append("")
lineas.append("## Conclusiones")
lineas.append("")
lineas.append("- [Análisis interpretativo basado en comparaciones]")
lineas.append("- [Recomendaciones específicas]")

reporte = '\\n'.join(lineas)

timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
output_file = '/tmp/analisis_' + timestamp + '.md'
with open(output_file, 'w', encoding='utf-8') as f:
    f.write(reporte)

print("======================================================================")
print("Reporte generado exitosamente")
print("======================================================================")
print("Archivo: " + output_file)
print("======================================================================")
cursor.close()
conn.close()
```

Genera el script siguiendo estas pautas de sintaxis.
"""

    # Intento 1: Usar RAGChain si está disponible
    if rag_chain:
        try:
            response = rag_chain.invoke(
                prompt,
                context_docs=[],
                use_llm=True
            )
            result = response['output'] if isinstance(response, dict) else str(response)
            if result and len(result) > 100:  # Validar que sea un resultado válido
                return result
        except Exception as e:
            import traceback
            error_msg = f"⚠️ RAGChain falló: {str(e)}\n{traceback.format_exc()}"
            print(error_msg)
            st.warning(f"RAGChain no disponible, usando Claude API directo...")

    # Intento 2: Fallback a Claude API directo
    try:
        import os
        api_key = os.getenv('ANTHROPIC_API_KEY')
        if not api_key:
            st.error("❌ ANTHROPIC_API_KEY no encontrada en variables de entorno")
            return None

        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model="claude-sonnet-4-20250514",  # Modelo Sonnet 4
            max_tokens=8192,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return message.content[0].text
    except Exception as e:
        import traceback
        error_msg = f"❌ Claude API falló: {str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        st.error(error_msg)
        return None


def execute_analysis_script(script_code: str) -> tuple[bool, str, str]:
    """
    Ejecuta el script Python generado y retorna:
    - success: bool
    - output: stdout del script
    - error: stderr si hay error
    """
    import subprocess
    import tempfile
    import re

    try:
        # Limpiar código: remover backticks de markdown si existen
        cleaned_code = script_code.strip()

        # Remover ```python al inicio y ``` al final
        if cleaned_code.startswith('```python'):
            cleaned_code = cleaned_code.split('```python', 1)[1]
        elif cleaned_code.startswith('```'):
            cleaned_code = cleaned_code.split('```', 1)[1]

        if cleaned_code.endswith('```'):
            cleaned_code = cleaned_code.rsplit('```', 1)[0]

        cleaned_code = cleaned_code.strip()

        # Validar sintaxis antes de ejecutar
        try:
            compile(cleaned_code, '<string>', 'exec')
        except SyntaxError as e:
            return False, "", f"Error de sintaxis en el código generado:\nLínea {e.lineno}: {e.msg}\n{e.text or ''}"

        # Guardar script en archivo temporal
        with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False) as f:
            f.write(cleaned_code)
            temp_script = f.name

        # Ejecutar script
        result = subprocess.run(
            ['python3', temp_script],
            capture_output=True,
            text=True,
            timeout=60
        )

        # Limpiar archivo temporal
        os.unlink(temp_script)

        if result.returncode == 0:
            return True, result.stdout, ""
        else:
            return False, result.stdout, result.stderr

    except subprocess.TimeoutExpired:
        return False, "", "El script tardó demasiado (timeout 60s)"
    except Exception as e:
        return False, "", str(e)


def read_markdown_report(output_text: str) -> str:
    """
    Extrae el path del archivo generado del output del script.
    Si no lo encuentra en el output, busca el archivo .md más reciente en /tmp/
    """
    import glob

    # Intentar extraer del output primero
    lines = output_text.split('\n')
    for line in lines:
        # Buscar cualquier archivo .md en /tmp/
        if '/tmp/' in line and '.md' in line:
            # Extraer path
            parts = line.split('/tmp/')
            if len(parts) > 1:
                # Tomar todo hasta encontrar un espacio, comilla o final de línea
                path_part = parts[1].split()[0].strip()
                # Remover caracteres no deseados al final
                path_part = path_part.rstrip('.,;:)\'"')
                full_path = '/tmp/' + path_part
                if os.path.exists(full_path):
                    return full_path

    # Si no se encontró en el output, buscar el archivo .md más reciente en /tmp/
    md_files = glob.glob('/tmp/*.md')
    if md_files:
        # Ordenar por tiempo de modificación, más reciente primero
        md_files.sort(key=os.path.getmtime, reverse=True)
        # Retornar el más reciente si fue modificado en los últimos 30 segundos
        most_recent = md_files[0]
        import time
        if time.time() - os.path.getmtime(most_recent) < 30:
            return most_recent

    return None


def convert_md_to_docx(md_path: str) -> str:
    """
    Convierte un archivo Markdown a DOCX
    Retorna el path del archivo DOCX generado
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Procesar línea por línea
    for line in md_content.split('\n'):
        line = line.strip()

        if not line:
            continue

        # Encabezados
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        # Separadores
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
        # Listas
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Negritas
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
        # Tablas (detección simple)
        elif '|' in line and not line.startswith('|---'):
            # Agregar como texto (parsing completo de tablas es más complejo)
            doc.add_paragraph(line)
        # Texto normal
        else:
            doc.add_paragraph(line)

    # Guardar
    docx_path = md_path.replace('.md', '.docx')
    doc.save(docx_path)
    return docx_path


def convert_md_to_pdf(md_path: str) -> str:
    """
    Convierte un archivo Markdown a PDF usando reportlab
    Retorna el path del archivo PDF generado
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    pdf_path = md_path.replace('.md', '.pdf')
    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Procesar línea por línea
    import re
    for line in md_content.split('\n'):
        line_original = line
        line = line.strip()

        if not line:
            story.append(Spacer(1, 0.2*inch))
            continue

        # Escapar caracteres especiales XML antes de procesamiento
        line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        # Procesar negritas **texto** → <b>texto</b>
        line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)

        # Encabezados
        if line_original.startswith('# '):
            story.append(Paragraph(line[2:], styles['Title']))
            story.append(Spacer(1, 0.3*inch))
        elif line_original.startswith('## '):
            story.append(Paragraph(line[3:], styles['Heading1']))
            story.append(Spacer(1, 0.2*inch))
        elif line_original.startswith('### '):
            story.append(Paragraph(line[4:], styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
        # Separadores
        elif line_original.startswith('---'):
            story.append(Spacer(1, 0.2*inch))
        # Listas
        elif line_original.startswith('- '):
            story.append(Paragraph('• ' + line[2:], styles['BodyText']))
        # Texto normal
        else:
            try:
                story.append(Paragraph(line, styles['BodyText']))
            except Exception as e:
                # Si falla el parseo, agregar texto plano
                safe_line = line.replace('<b>', '').replace('</b>', '')
                story.append(Paragraph(safe_line, styles['BodyText']))

    # Construir PDF
    doc.build(story)
    return pdf_path


# ============ INTERFAZ PRINCIPAL ============

def app():
    # ============ INICIALIZAR SESSION STATE ============
    if 'historial_scripts' not in st.session_state:
        st.session_state.historial_scripts = []
    if 'script_a_ejecutar' not in st.session_state:
        st.session_state.script_a_ejecutar = None
    if 'peticion_original' not in st.session_state:
        st.session_state.peticion_original = None

    col1, col2 = st.columns([3, 1])
    with col1:
        st.markdown('<div class="main-header">🤖 Generador de Análisis Inteligente</div>',
                   unsafe_allow_html=True)

    with col2:
        st.markdown("**LATAM-3C**")

    st.markdown("""
    Describe el análisis que necesitas y la IA generará un script Python, lo ejecutará
    y mostrará el reporte automáticamente.
    """)

    # ============ HISTORIAL DE SCRIPTS ============
    if len(st.session_state.historial_scripts) > 0:
        with st.expander(f"📚 Historial de Scripts ({len(st.session_state.historial_scripts)})", expanded=False):
            st.markdown("**Selecciona un script previo para ejecutarlo nuevamente:**")

            for idx, item in enumerate(reversed(st.session_state.historial_scripts[-10:])):  # Últimos 10
                col_a, col_b = st.columns([4, 1])
                with col_a:
                    st.markdown(f"**{idx + 1}.** {item['peticion'][:70]}...")
                    st.caption(f"🕒 {item['timestamp']}")
                with col_b:
                    if st.button("▶️ Ejecutar", key=f"exec_hist_{idx}", use_container_width=True):
                        st.session_state.script_a_ejecutar = item['script']
                        st.session_state.peticion_original = item['peticion']
                        st.rerun()
                if idx < len(st.session_state.historial_scripts[-10:]) - 1:
                    st.divider()

    st.divider()

    # ============ SECCIÓN DE ENTRADA ============
    st.markdown('<div class="section-header">📝 Tu Petición de Análisis</div>',
               unsafe_allow_html=True)

    # Usar ejemplo seleccionado como valor inicial
    valor_inicial = ""
    if 'ejemplo_seleccionado' in st.session_state and st.session_state.ejemplo_seleccionado:
        valor_inicial = st.session_state.ejemplo_seleccionado
        st.session_state.ejemplo_seleccionado = None  # Limpiar después de usar

    col1, col2 = st.columns([3, 1])
    with col1:
        peticion = st.text_area(
            "Escribe qué análisis necesitas",
            value=valor_inicial,
            placeholder="Ej: Analiza la distribución de huella CO₂ por empresa y año, mostrando gráficos de tendencias",
            height=100,
            key="peticion_input"
        )

    with col2:
        st.markdown("")
        st.markdown("")
        generar = st.button("🚀 Generar Análisis", use_container_width=True)

    # ============ PROCESAMIENTO DESDE HISTORIAL ============
    if st.session_state.script_a_ejecutar:
        script_code = st.session_state.script_a_ejecutar
        peticion = st.session_state.peticion_original

        st.info(f"📌 Ejecutando script del historial: {peticion[:100]}...")

        # Limpiar flags
        st.session_state.script_a_ejecutar = None
        st.session_state.peticion_original = None

        # Mostrar script
        with st.expander("👁️ Ver script"):
            st.code(script_code, language="python")

        # Ejecutar directamente (saltar a paso 2)
        st.markdown('<div class="section-header">⚙️ Ejecutando Script</div>',
                   unsafe_allow_html=True)

        with st.spinner("⏳ Ejecutando análisis..."):
            success, output, error = execute_analysis_script(script_code)

        if success:
            st.success("✅ Script ejecutado exitosamente")
            report_path = read_markdown_report(output)

            if report_path and os.path.exists(report_path):
                st.markdown('<div class="section-header">📊 Reporte</div>',
                           unsafe_allow_html=True)

                with open(report_path, 'r', encoding='utf-8') as f:
                    reporte_md = f.read()

                with st.expander("📄 Ver Reporte Completo", expanded=True):
                    st.markdown(reporte_md)

                # Botones de descarga en múltiples formatos
                col1, col2, col3 = st.columns(3)

                with col1:
                    with open(report_path, 'rb') as f:
                        st.download_button(
                            label="📥 Descargar MD",
                            data=f,
                            file_name=os.path.basename(report_path),
                            mime='text/markdown',
                            use_container_width=True
                        )

                with col2:
                    try:
                        docx_path = convert_md_to_docx(report_path)
                        with open(docx_path, 'rb') as f:
                            st.download_button(
                                label="📄 Descargar DOCX",
                                data=f,
                                file_name=os.path.basename(docx_path),
                                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                use_container_width=True
                            )
                    except Exception as e:
                        st.error(f"Error al generar DOCX: {str(e)}")

                with col3:
                    try:
                        pdf_path = convert_md_to_pdf(report_path)
                        with open(pdf_path, 'rb') as f:
                            st.download_button(
                                label="📕 Descargar PDF",
                                data=f,
                                file_name=os.path.basename(pdf_path),
                                mime='application/pdf',
                                use_container_width=True
                            )
                    except Exception as e:
                        st.error(f"Error al generar PDF: {str(e)}")
            else:
                st.warning("⚠️ No se encontró archivo de reporte")
                if output:
                    with st.expander("Ver output del script"):
                        st.code(output)
        else:
            st.error("❌ Error al ejecutar el script")
            if error:
                with st.expander("🔍 Ver detalles del error", expanded=True):
                    st.code(error, language="text")

        return  # Terminar aquí para no continuar con el flujo normal

    # ============ PROCESAMIENTO GENERACIÓN NUEVA ============
    if generar and peticion.strip():

        # Inicializar RAGChain
        with st.spinner("⏳ Inicializando..."):
            try:
                rag_chain = RAGChain()
            except Exception as e:
                st.error(f"❌ Error al inicializar: {str(e)}")
                return

        # Paso 1: Generar script
        st.markdown('<div class="section-header">📋 Paso 1: Generando Script</div>',
                   unsafe_allow_html=True)

        with st.spinner("🤖 Generando script Python..."):
            script_code = generate_analysis_script(peticion, rag_chain)

            if not script_code:
                st.error("❌ Error al generar el script. Intenta de nuevo.")
                return

        st.success("✅ Script generado exitosamente")

        # Guardar en historial
        from datetime import datetime
        st.session_state.historial_scripts.append({
            'peticion': peticion,
            'script': script_code,
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })
        # Limitar a últimos 10 scripts
        if len(st.session_state.historial_scripts) > 10:
            st.session_state.historial_scripts = st.session_state.historial_scripts[-10:]

        # Mostrar script en expander
        with st.expander("👁️ Ver script generado"):
            st.code(script_code, language="python")

        # Paso 2: Ejecutar script
        st.markdown('<div class="section-header">⚙️ Paso 2: Ejecutando Script</div>',
                   unsafe_allow_html=True)

        with st.spinner("⏳ Ejecutando análisis..."):
            success, output, error = execute_analysis_script(script_code)

        if success:
            st.success("✅ Script ejecutado exitosamente")

            # Extraer path del reporte
            report_path = read_markdown_report(output)

            if report_path and os.path.exists(report_path):
                st.markdown('<div class="section-header">📊 Paso 3: Reporte</div>',
                           unsafe_allow_html=True)

                # Leer y mostrar reporte en expander
                with open(report_path, 'r', encoding='utf-8') as f:
                    reporte_md = f.read()

                with st.expander("📄 Ver Reporte Completo", expanded=True):
                    st.markdown(reporte_md)

                # Botones de descarga en múltiples formatos
                st.divider()
                col1, col2, col3 = st.columns(3)

                with col1:
                    with open(report_path, 'rb') as f:
                        st.download_button(
                            label="📥 Descargar MD",
                            data=f,
                            file_name=os.path.basename(report_path),
                            mime='text/markdown',
                            use_container_width=True
                        )

                with col2:
                    docx_path = convert_md_to_docx(report_path)
                    with open(docx_path, 'rb') as f:
                        st.download_button(
                            label="📄 Descargar DOCX",
                            data=f,
                            file_name=os.path.basename(docx_path),
                            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                            use_container_width=True
                        )

                with col3:
                    pdf_path = convert_md_to_pdf(report_path)
                    with open(pdf_path, 'rb') as f:
                        st.download_button(
                            label="📕 Descargar PDF",
                            data=f,
                            file_name=os.path.basename(pdf_path),
                            mime='application/pdf',
                            use_container_width=True
                        )
            else:
                st.warning("⚠️ No se encontró el archivo de reporte generado")
                if output:
                    with st.expander("🔍 Ver output del script", expanded=False):
                        st.code(output)
        else:
            st.error("❌ Error al ejecutar el script")
            if error:
                with st.expander("🔍 Ver detalles del error", expanded=True):
                    st.code(error, language="text")
            if output:
                with st.expander("📋 Ver output del script", expanded=False):
                    st.code(output)

    elif generar and not peticion.strip():
        st.warning("⚠️ Por favor escribe una petición de análisis")

    # ============ SECCIÓN DE EJEMPLOS ============
    st.divider()
    st.markdown('<div class="section-header">💡 Ejemplos de Peticiones</div>',
               unsafe_allow_html=True)

    ejemplos = [
        "Analiza la tendencia de huella CO₂ por empresa (MZMA, Melón, LOMAX, Pacas) desde 2020 a 2024",
        "Compara la resistencia (REST) promedio entre empresas y crea una tabla de clasificación",
        "Identifica los outliers en huella de carbono (valores fuera del rango 150-450 kg CO₂/m³)",
        "Analiza la distribución de volumen producido por año y empresa",
        "Calcula estadísticas por tipo de cemento usando la tabla ref_tipos_cemento"
    ]

    # Inicializar ejemplo seleccionado en session_state
    if 'ejemplo_seleccionado' not in st.session_state:
        st.session_state.ejemplo_seleccionado = None

    cols = st.columns(1)
    for i, ejemplo in enumerate(ejemplos):
        with cols[0]:
            if st.button(f"📌 {ejemplo}", use_container_width=True, key=f"ejemplo_{i}"):
                st.session_state.ejemplo_seleccionado = ejemplo
                st.rerun()

    # ============ INFORMACIÓN ============
    st.divider()
    st.markdown("""
    ### ℹ️ Cómo funciona

    1. **Describe tu análisis**: Usa lenguaje natural en español
    2. **IA genera script**: Claude crea un script Python optimizado para PostgreSQL
    3. **Ejecución automática**: El script se ejecuta en tu servidor
    4. **Reporte interactivo**: Ves el resultado formateado en Markdown
    5. **Descarga**: Puedes guardar el reporte en tu computadora

    ### 📊 Datos disponibles
    - **260 registros** de huella de concreto (2020-2024)
    - **4 empresas** principales: MZMA, Melón, LOMAX, Pacas
    - **139 tipos de cemento** con factor clinker
    - **265 plantas** en LATAM
    """)


if __name__ == "__main__":
    app()
